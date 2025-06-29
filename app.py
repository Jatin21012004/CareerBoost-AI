import sys
import streamlit as st
from resume_parser import extract_text_from_pdf, extract_text_from_docx, parse_resume
from job_matcher import match_resume_to_job, extract_skills_from_text, get_skill_gaps
from suggestions import generate_suggestions
import io
from docx import Document
from fpdf import FPDF
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image
import pytesseract
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from datetime import datetime
from pymongo import MongoClient
import os
from dotenv import load_dotenv
import openai
from huggingface_hub import InferenceApi

# ========== Load environment variables ==========
load_dotenv()
api_token = os.getenv("HUGGINGFACE_API_KEY")
api = InferenceApi(repo_id="mistralai/Mistral-7B-Instruct-v0.1", token=api_token)

# ========== AI Career Coach Chat ==========
# ========== AI Career Coach Chat (Mistral Style) ==========
def career_coach_chat(user_query):
    try:
        # Format input for Mistral Instruct
        prompt = f"[INST] {user_query.strip()} [/INST]"
        
        # Send to Hugging Face API
        response = api(inputs=prompt, raw_response=True)
        response_json = response.json()
        
        # Extract the text (check multiple keys)
        generated_text = response_json.get("generated_text", "").strip()
        
        # If Mistral responds with 'text' inside a list
        if not generated_text and isinstance(response_json, list):
            generated_text = response_json[0].get("generated_text", "").strip()

        # If no valid text received
        if not generated_text:
            return "💡 I couldn't generate a reply right now. Please try rephrasing or asking another question."

        return generated_text

    except Exception as e:
        return f"🚨 Error: {str(e)}"


# ========== Image Text Extraction ==========
def extract_text_from_image(uploaded_file):
    try:
        image = Image.open(uploaded_file)
        return pytesseract.image_to_string(image)
    except Exception as e:
        st.error(f"OCR Error: {str(e)}")
        return ""

# ========== Generate PDF Report ==========
def generate_pdf_report(suggestions, resume_data, match_score):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.setFont("Helvetica", 16)
    p.drawString(100, 750, "Resume Improvement Report")
    p.setFont("Helvetica", 12)
    text = p.beginText(100, 700)
    text.textLine(f"Candidate: {resume_data.get('name', 'N/A')}")
    text.textLine(f"Match Score: {match_score}%")
    text.textLine("")
    text.textLine("Key Suggestions:")
    text.textLine("")
    for s in suggestions:
        clean_suggestion = s.replace('•', '-').replace('**','').strip()
        text.textLine(f"- {clean_suggestion}")
    p.drawText(text)
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer.read()

# ========== Generate DOCX Report ==========
def generate_docx_report(suggestions, resume_data, match_score):
    doc = Document()
    doc.add_heading("Resume Improvement Report", 0)
    doc.add_paragraph(f"Candidate: {resume_data.get('name', 'N/A')}")
    doc.add_paragraph(f"Match Score: {match_score}%")
    doc.add_heading("Key Suggestions", level=1)
    for s in suggestions:
        doc.add_paragraph(s.strip(), style='List Bullet')
    doc_io = io.BytesIO()
    doc.save(doc_io)
    return doc_io.getvalue()

# ========== Streamlit App Setup ==========
st.set_page_config(
    page_title="AI Resume Analyzer Pro",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="expanded"
)

# ========== Sidebar: AI Career Coach ==========
with st.sidebar:
    st.title("🤖 Resume Analyzer Pro")

    if "chat_messages" not in st.session_state:
        st.session_state.chat_messages = []

    with st.expander("💬 AI Career Coach", expanded=True):
        user_query = st.text_input("Ask career questions...", key="chat_input")

        if user_query:
            if any(greeting in user_query.lower() for greeting in ["hi", "hello", "hey"]):
                response = "Hello! How can I assist you with your career today?"
            else:
                response = career_coach_chat(user_query)

            st.session_state.chat_messages.append({"role": "user", "content": user_query})
            st.session_state.chat_messages.append({"role": "assistant", "content": response})

        for msg in st.session_state.chat_messages[-6:]:
            st.chat_message(msg["role"]).write(msg["content"])

    st.markdown("---")
    st.markdown("Made with ❤️ by **Jatin Jawa & Team**")
    st.markdown("### How to Use:")
    st.markdown("1. Upload your resume (PDF/DOCX/Image)")
    st.markdown("2. Paste the job description")
    st.markdown("3. Get AI-powered feedback!")
    st.markdown("---")
    st.markdown("🚀 *Empowering Job Seekers with AI*")

# ========== Main Interface ==========
st.title("📄 AI-Powered Resume Analyzer")
st.caption("Optimize your resume for any job description in seconds")

col1, col2 = st.columns(2)
with col1:
    uploaded_file = st.file_uploader(
        "Upload Resume (PDF/DOCX/Image)",
        type=["pdf", "docx", "png", "jpg", "jpeg"],
        help="We support scanned documents too!"
    )
with col2:
    job_description = st.text_area(
        "Paste Job Description",
        height=150,
        placeholder="Paste the full job description here..."
    )

if uploaded_file and job_description:
    with st.spinner("🔍 Analyzing your resume..."):
        file_text = ""
        if uploaded_file.name.endswith(".pdf"):
            file_text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.endswith(".docx"):
            file_text = extract_text_from_docx(uploaded_file)
        else:
            file_text = extract_text_from_image(uploaded_file)

        if not file_text.strip():
            st.error("⚠️ Could not extract text from file. Try a different format.")
            st.stop()

        resume_data = parse_resume(file_text)
        match_score = match_resume_to_job(file_text, job_description)
        suggestions = generate_suggestions(resume_data, match_score, job_description)

        pdf_bytes = generate_pdf_report(suggestions, resume_data, match_score)
        docx_bytes = generate_docx_report(suggestions, resume_data, match_score)

    st.success("✅ Analysis Complete!")
    st.subheader("🚀 Personalized Suggestions")
    for suggestion in suggestions:
        if "🔴" in suggestion:
            st.error(suggestion, icon="❌")
        elif "🟡" in suggestion:
            st.warning(suggestion, icon="⚠️")
        elif "🟢" in suggestion:
            st.success(suggestion, icon="✅")
        else:
            st.info(suggestion, icon="ℹ️")

    st.download_button("📥 Download Suggestions as PDF", data=pdf_bytes, file_name="resume_suggestions.pdf", key="download_pdf_1")
    st.download_button("📥 Download Suggestions as DOCX", data=docx_bytes, file_name="resume_suggestions.docx", key="download_docx_1")

    col_a, col_b = st.columns(2)
    with col_a:
        st.metric("Overall Match Score", f"{match_score}%")
    with col_b:
        if match_score < 40:
            st.error("🔴 Needs Major Improvement")
        elif match_score < 70:
            st.warning("🟡 Good But Could Improve")
        else:
            st.success("🟢 Excellent Match!")

    st.divider()

    with st.expander("🔍 View Detailed Analysis", expanded=False):
        tab1, tab2, tab3 = st.tabs(["Resume Data", "Skill Matching", "Raw Text"])

        with tab1:
            st.subheader("Extracted Resume Data")
            st.json(resume_data)

        with tab2:
            st.subheader("Skill Analysis")
            resume_skills = resume_data.get("skills", [])
            job_skills = extract_skills_from_text(job_description)

            col_x, col_y = st.columns(2)
            with col_x:
                with st.container(border=True):
                    st.metric("Your Skills", len(resume_skills))
                    st.write(", ".join(resume_skills) if resume_skills else "None found")

            with col_y:
                with st.container(border=True):
                    st.metric("Required Skills", len(job_skills))
                    st.write(", ".join(job_skills) if job_skills else "None specified")

            if missing_skills := get_skill_gaps(resume_skills, job_skills):
                st.error(f"Missing {len(missing_skills)} Key Skills:")
                st.write(", ".join(missing_skills.keys()))

        with tab3:
            st.subheader("Raw Resume Text")
            st.code(file_text[:5000] + "..." if len(file_text) > 5000 else file_text)

elif uploaded_file and not job_description:
    st.error("Please paste the job description to analyze!")
elif not uploaded_file and job_description:
    st.error("Please upload your resume first!")
else:
    st.info("👆 Upload your resume and paste a job description to begin")
